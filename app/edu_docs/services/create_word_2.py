import os
import re
import uuid
import lxml

from django.conf import settings
from django.core.exceptions import ImproperlyConfigured

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
import datetime
import docx

def get_dated_path(filename):
    today = datetime.datetime.now()
    return os.path.join('documents', today.strftime('%Y'), today.strftime('%m'), today.strftime('%d'), filename)


def save_doc_in_media(doc, sanitized_theme):
    try:
        sanitized_theme = f'{sanitized_theme}.docx'
        # Создание пути с текущей датой
        dated_path = get_dated_path(sanitized_theme)
        full_path = os.path.join(settings.MEDIA_ROOT, dated_path)

        # Создание всех необходимых директорий
        os.makedirs(os.path.dirname(full_path), exist_ok=True)

        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        element_updatefields = lxml.etree.SubElement(
            doc.settings.element, f"{namespace}updateFields"
        )
        element_updatefields.set(f"{namespace}val", "true")

        doc.save(full_path)
        return full_path, sanitized_theme  # Возвращаем относительный путь

    except OSError as e:
        # Обработка ошибок создания директории или сохранения файла
        raise ImproperlyConfigured(f"Failed to save document: {e}")
    except Exception as e:
        # Обработка других исключений
        raise Exception(f"An unexpected error occurred: {e}")


def sanitize_filename(filename):
    """
        Функция sanitize_filename удаляет из названия файла символы, которые Word не допускает.

        Принимает:
        - filename (str): исходное название файла.

        Возвращает:
        - sanitized_filename (str): отредактированное название файла.
    """

    # Заменяем запрещенные символы на пустую строку
    sanitized_filename = re.sub(r'[\\/:\*\?"<>|+%!.@]', '', filename)

    # Убираем пробелы и точки в конце имени файла
    sanitized_filename = sanitized_filename.rstrip(' .')

    return sanitized_filename


def create_word(work_theme, subtopics, texts, university, work_type, author_name, group_name, teacher_name, language_of_work, cover_page_data):
    """
        Функция createword оформляет и сохраняет Word документ по стандарту.

        Принимает:
        - work_theme (str): тема работы.
        - subtopics (list): массив подглав.
        - texts (list): массив текстов для подглав.
        - university (str): название университета.
        - work_type (str): тип работы.
        - author_name (str): имя автора.
        - group_name (str): название группы.
        - teacher_name (str): имя преподавателя.
        - language_of_work (str): язык работы.
        - cover_page_data (str): данные для титульного листа.

        Функция создает Word документ, форматирует его по стандарту и сохраняет в тот же каталог, где находится основной файл (main).
    """

    # Заменяем запрещенные символы на пустую строку в название файла
    sanitized_theme = sanitize_filename(work_theme)

    # Если язык работы кыргызский
    if language_of_work == "1":
        DEFAULT_FONT_NAME = "Times New Roman"

        # Функция для создания элемента
        def create_element(name):
            return OxmlElement(name)

        # Функция для создания атрибута
        def create_attribute(element, name, value):
            element.set(ns.qn(name), value)

        # Функция для добавления номера страницы
        def add_page_number(section):
            p = section.footer.paragraphs[0]  # Получаем параграф из нижнего колонтитула
            p.alignment = 1  # Выравнивание по центру
            p.style.font.size = Pt(12)  # Размер шрифта
            p.style.font.name = DEFAULT_FONT_NAME  # Шрифт

            # Создаем элементы XML для номера страницы
            run = p.add_run()
            fldChar1 = create_element('w:fldChar')
            create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = create_element('w:instrText')
            create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = create_element('w:fldChar')
            create_attribute(fldChar2, 'w:fldCharType', 'end')

            # Добавляем элементы в параграф для форматирования номера страницы
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        def add_heading(document, text, level=1):
            # Добавляем заголовок
            heading = document.add_heading(text, level=level)
            heading.bold = True
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.text = run.text.upper()  # Преобразуем текст заголовка в верхний регистр
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if text == "Киришүү" or text == "Корутунду" or text == "Колдонулган адабияттар":
                heading.paragraph_format.line_spacing = 2  # Задаем межстрочный интервал
            else:
                heading.paragraph_format.line_spacing = 1  # Задаем межстрочный интервал


        def add_subheading(document, text, level=2):
            # Добавляем подзаголовок
            heading = document.add_heading(level=level)
            heading_run = heading.add_run(text)
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.paragraph_format.line_spacing = 2  # Задаем межстрочный интервал

        def add_text(document, text):
            # Добавляем текст
            style = document.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            paragraphs = text.split("\n")
            for paragraph_text in paragraphs:
                paragraph = document.add_paragraph()
                paragraph.add_run("\t" + paragraph_text)  # Добавляем текст с отступом
                paragraph.paragraph_format.right_indent = docx.shared.Inches(-0.590551)  # Устанавливаем правый отступ


        doc = Document()

        doc.sections[0].top_margin = docx.shared.Cm(2)
        # Если надо добавить титульный лист
        if cover_page_data == "1" or cover_page_data == "3":
            # Добавляем текст "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КЫРГЫЗСКОЙ РЕСПУБЛИКИ"
            ministry_title = doc.add_paragraph("КЫРГЫЗ РЕСПУБЛИКАСЫНЫН БИЛИМ БЕРҮҮ ЖАНА ИЛИМ МИНИСТИРЛИГИ")
            ministry_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            ministry_title.runs[0].font.name = "Times New Roman"
            ministry_title.runs[0].font.size = Pt(14)
            ministry_title.paragraph_format.line_spacing = 1.15

            # Добавляем название университета
            uni_title = doc.add_paragraph(university)
            uni_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            uni_title.runs[0].font.name = "Times New Roman"
            uni_title.runs[0].font.size = Pt(14)
            uni_title.runs[0].text = uni_title.runs[0].text.upper()
            uni_title.paragraph_format.line_spacing = 1.15

            if work_type == "Студенттин өз алдынча иши":
                # Добавляем пустые строки
                for _ in range(3):
                    doc.add_paragraph()
            else:
                # Добавляем пустые строки
                for _ in range(4):
                    doc.add_paragraph()

            # Добавляем тип работы
            type_works_title = doc.add_paragraph(work_type)
            type_works_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            type_works_title.runs[0].font.name = "Times New Roman"
            type_works_title.runs[0].font.size = Pt(36)
            type_works_title.runs[0].bold = True
            type_works_title.paragraph_format.line_spacing = 1.15

            # Добавляем название темы
            theme_title = doc.add_paragraph("Тема: " + work_theme)
            theme_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            theme_title.runs[0].font.name = "Times New Roman"
            theme_title.runs[0].font.size = Pt(14)
            theme_title.paragraph_format.line_spacing = 1.15

            if work_type == "Студенттин өз алдынча иши":
                # Проверяем длину темы
                if len(work_theme) < 59:
                    # Добавляем пустые строки
                    for _ in range(4):
                        doc.add_paragraph()
                elif 59 <= len(work_theme) < 113:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 114 <= len(work_theme) < 179:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 180 <= len(work_theme):
                    # Добавляем пустые строки
                    for _ in range(2):
                        doc.add_paragraph()
            else:
                # Проверяем длину темы
                if len(work_theme) < 113:
                    # Добавляем пустые строки
                    for _ in range(4):
                        doc.add_paragraph()
                elif 114 <= len(work_theme) < 179:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 180 <= len(work_theme):
                    # Добавляем пустые строки
                    for _ in range(1):
                        doc.add_paragraph()


            # Добавляем текст "Выполнил(а):" ФИО
            author_title = doc.add_paragraph("Аткарган: " + author_name)
            author_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            author_title.runs[0].font.name = "Times New Roman"
            author_title.runs[0].font.size = Pt(14)
            author_title.paragraph_format.line_spacing = 1.15

            # Добавляем текст "Группа:" название группы
            group_title = doc.add_paragraph("Группа: " + group_name)
            group_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            group_title.runs[0].font.name = "Times New Roman"
            group_title.runs[0].font.size = Pt(14)
            group_title.paragraph_format.line_spacing = 1.15

            # Добавляем текст "Проверил(а):" ФИО преподавателя
            teacher_title = doc.add_paragraph("Текшерген: " + teacher_name)
            teacher_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            teacher_title.runs[0].font.name = "Times New Roman"
            teacher_title.runs[0].font.size = Pt(14)
            teacher_title.paragraph_format.line_spacing = 1.15

            if len(university) <= 89:
                # Добавляем пустые строки
                for _ in range(3):
                    doc.add_paragraph()
            elif len(university) > 89:
                # Добавляем пустые строки
                for _ in range(2):
                    doc.add_paragraph()

            # Получаем текущую дату и извлекаем год
            current_year = datetime.datetime.now().year

            # Добавляем текст "Бишкек -" текущий год
            centered_paragraph = doc.add_paragraph()
            centered_run = centered_paragraph.add_run("Бишкек - " + str(current_year))
            centered_run.font.name = "Times New Roman"
            centered_run.font.size = Pt(14)
            centered_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            centered_paragraph.paragraph_format.line_spacing = 1.15

            doc.add_paragraph()

        # Добавляем заголовок для оглавления
        toc_title = doc.add_paragraph("МАЗМУНУ")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in toc_title.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True

        # Добавляем пустой параграф перед оглавлением
        toc_empty_para = doc.add_paragraph()
        toc_empty_para.paragraph_format.space_before = Inches(0)
        toc_empty_para.paragraph_format.space_after = Inches(0)
        toc_run = toc_empty_para.add_run()

        # Создаем XML-элементы для оглавления
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # измените "1-3" в зависимости от необходимого уровня заголовков

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar3 = OxmlElement('w:instrText')  # Заменяем на 'w:instrText', чтобы добавить текст инструкции
        fldChar3.set(qn('xml:space'), 'preserve')
        fldChar3.text = "Right-click to update field."

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        # Добавляем XML-элементы в оглавление
        toc_run._r.append(fldChar)
        toc_run._r.append(instrText)
        toc_run._r.append(fldChar2)
        toc_run._r.append(fldChar3)  # Добавляем элемент с текстом инструкции
        toc_run._r.append(fldChar4)

        # Добавляем разрыв страницы
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Добавляем раздел, из которого мы хотим получить номера страниц
        doc.add_section(WD_SECTION.CONTINUOUS)

        # Добавляем текст на первой странице
        doc.sections[1].footer.is_linked_to_previous = False

        add_heading(doc, subtopics[0])
        add_text(doc, texts[0])
        doc.add_page_break()

        add_heading(doc, "1-БӨЛҮМ. " + subtopics[1])
        add_subheading(doc, "1.1. " + subtopics[2])
        add_text(doc, texts[1])
        doc.add_page_break()

        add_subheading(doc, "1.2. " + subtopics[3])
        add_text(doc, texts[2])
        doc.add_page_break()

        add_subheading(doc, "1.3. " + subtopics[4])
        add_text(doc, texts[3])
        doc.add_page_break()

        add_heading(doc, "2-БӨЛҮМ. " + subtopics[5])
        add_subheading(doc, "2.1. " + subtopics[6])
        add_text(doc, texts[4])
        doc.add_page_break()

        add_subheading(doc, "2.2. " + subtopics[7])
        add_text(doc, texts[5])
        doc.add_page_break()

        add_subheading(doc, "2.3. " + subtopics[8])
        add_text(doc, texts[6])
        doc.add_page_break()

        add_heading(doc, subtopics[9])
        add_text(doc, texts[7])
        doc.add_page_break()

        add_heading(doc, subtopics[10])
        bibliography = texts[8].split("\n")
        for i, reference in enumerate(bibliography, start=1):
            doc.add_paragraph(f"{i}. {reference.strip()}")

        # Добавляем номера страниц
        add_page_number(doc.sections[1])

        return save_doc_in_media(doc, sanitized_theme)


    # Если язык работы русский
    if language_of_work == "2":
        DEFAULT_FONT_NAME = "Times New Roman"

        # Функция для создания элемента
        def create_element(name):
            return OxmlElement(name)

        # Функция для создания атрибута
        def create_attribute(element, name, value):
            element.set(ns.qn(name), value)

        # Функция для добавления номера страницы
        def add_page_number(section):
            p = section.footer.paragraphs[0]  # Получаем параграф из нижнего колонтитула
            p.alignment = 1  # Выравнивание по центру
            p.style.font.size = Pt(12)  # Размер шрифта
            p.style.font.name = DEFAULT_FONT_NAME  # Шрифт

            # Создаем элементы XML для номера страницы
            run = p.add_run()
            fldChar1 = create_element('w:fldChar')
            create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = create_element('w:instrText')
            create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = create_element('w:fldChar')
            create_attribute(fldChar2, 'w:fldCharType', 'end')

            # Добавляем элементы в параграф для форматирования номера страницы
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        def add_heading(document, text, level=1):
            # Добавляем заголовок
            heading = document.add_heading(text, level=level)
            heading.bold = True
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.text = run.text.upper()  # Преобразуем текст заголовка в верхний регистр
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if text == "Введение" or text == "Заключение" or text == "Список литературы":
                heading.paragraph_format.line_spacing = 2  # Задаем межстрочный интервал
            else:
                heading.paragraph_format.line_spacing = 1  # Задаем межстрочный интервал

        def add_subheading(document, text, level=2):
            # Добавляем подзаголовок
            heading = document.add_heading(level=level)
            heading_run = heading.add_run(text)
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.paragraph_format.line_spacing = 2  # Задаем межстрочный интервал

        def add_text(document, text):
            # Добавляем текст
            style = document.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            paragraphs = text.split("\n")
            for paragraph_text in paragraphs:
                paragraph = document.add_paragraph()
                paragraph.add_run("\t" + paragraph_text)  # Добавляем текст с отступом
                paragraph.paragraph_format.right_indent = docx.shared.Inches(-0.590551)  # Устанавливаем правый отступ

        doc = Document()

        doc.sections[0].top_margin = docx.shared.Cm(2)
        # Если надо добавить титульный лист
        if cover_page_data == "1" or cover_page_data == "3":
            # Добавляем текст "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КЫРГЫЗСКОЙ РЕСПУБЛИКИ"
            ministry_title = doc.add_paragraph("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КЫРГЫЗСКОЙ РЕСПУБЛИКИ")
            ministry_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            ministry_title.runs[0].font.name = "Times New Roman"
            ministry_title.runs[0].font.size = Pt(14)
            ministry_title.paragraph_format.line_spacing = 1.15

            # Добавляем название университета
            uni_title = doc.add_paragraph(university)
            uni_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            uni_title.runs[0].font.name = "Times New Roman"
            uni_title.runs[0].font.size = Pt(14)
            uni_title.runs[0].text = uni_title.runs[0].text.upper()
            uni_title.paragraph_format.line_spacing = 1.15

            if work_type == "Самостоятельная работа студента":
                # Добавляем пустые строки
                for _ in range(3):
                    doc.add_paragraph()
            else:
                # Добавляем пустые строки
                for _ in range(4):
                    doc.add_paragraph()

            # Добавляем тип работы
            type_works_title = doc.add_paragraph(work_type)
            type_works_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            type_works_title.runs[0].font.name = "Times New Roman"
            type_works_title.runs[0].font.size = Pt(36)
            type_works_title.runs[0].bold = True
            type_works_title.paragraph_format.line_spacing = 1.15

            # Добавляем название темы
            theme_title = doc.add_paragraph("На тему: " + work_theme)
            theme_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            theme_title.runs[0].font.name = "Times New Roman"
            theme_title.runs[0].font.size = Pt(14)
            theme_title.paragraph_format.line_spacing = 1.15

            if work_type == "Самостоятельная работа студента":
                # Проверяем длину темы
                if len(work_theme) < 59:
                    # Добавляем пустые строки
                    for _ in range(4):
                        doc.add_paragraph()
                elif 59 <= len(work_theme) < 113:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 114 <= len(work_theme) < 179:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 180 <= len(work_theme):
                    # Добавляем пустые строки
                    for _ in range(2):
                        doc.add_paragraph()
            else:
                # Проверяем длину темы
                if len(work_theme) < 113:
                    # Добавляем пустые строки
                    for _ in range(4):
                        doc.add_paragraph()
                elif 114 <= len(work_theme) < 179:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 180 <= len(work_theme):
                    # Добавляем пустые строки
                    for _ in range(1):
                        doc.add_paragraph()

            # Добавляем текст "Выполнил(а):" ФИО
            author_title = doc.add_paragraph("Выполнил(а): " + author_name)
            author_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            author_title.runs[0].font.name = "Times New Roman"
            author_title.runs[0].font.size = Pt(14)
            author_title.paragraph_format.line_spacing = 1.15

            # Добавляем текст "Группа:" название группы
            group_title = doc.add_paragraph("Группа: " + group_name)
            group_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            group_title.runs[0].font.name = "Times New Roman"
            group_title.runs[0].font.size = Pt(14)
            group_title.paragraph_format.line_spacing = 1.15

            # Добавляем текст "Проверил(а):" ФИО преподавателя
            teacher_title = doc.add_paragraph("Проверил(а): " + teacher_name)
            teacher_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            teacher_title.runs[0].font.name = "Times New Roman"
            teacher_title.runs[0].font.size = Pt(14)
            teacher_title.paragraph_format.line_spacing = 1.15

            if len(university) <= 89:
                # Добавляем пустые строки
                for _ in range(3):
                    doc.add_paragraph()
            elif len(university) > 89:
                # Добавляем пустые строки
                for _ in range(2):
                    doc.add_paragraph()

            # Получаем текущую дату и извлекаем год
            current_year = datetime.datetime.now().year

            # Добавляем текст "Бишкек -" текущий год
            centered_paragraph = doc.add_paragraph()
            centered_run = centered_paragraph.add_run("Бишкек - " + str(current_year))
            centered_run.font.name = "Times New Roman"
            centered_run.font.size = Pt(14)
            centered_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            centered_paragraph.paragraph_format.line_spacing = 1.15

            doc.add_paragraph()

        # Добавляем заголовок для оглавления
        toc_title = doc.add_paragraph("ОГЛАВЛЕНИЕ")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in toc_title.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True

        # Добавляем пустой параграф перед оглавлением
        toc_empty_para = doc.add_paragraph()
        toc_empty_para.paragraph_format.space_before = Inches(0)
        toc_empty_para.paragraph_format.space_after = Inches(0)
        toc_run = toc_empty_para.add_run()

        # Создаем XML-элементы для оглавления
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # измените "1-3" в зависимости от необходимого уровня заголовков

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar3 = OxmlElement('w:instrText')  # Заменяем на 'w:instrText', чтобы добавить текст инструкции
        fldChar3.set(qn('xml:space'), 'preserve')
        fldChar3.text = "Right-click to update field."

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        # Добавляем XML-элементы в оглавление
        toc_run._r.append(fldChar)
        toc_run._r.append(instrText)
        toc_run._r.append(fldChar2)
        toc_run._r.append(fldChar3)  # Добавляем элемент с текстом инструкции
        toc_run._r.append(fldChar4)

        # Добавляем разрыв страницы
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Добавляем раздел, из которого мы хотим получить номера страниц
        doc.add_section(WD_SECTION.CONTINUOUS)

        # Добавляем текст на первой странице
        doc.sections[1].footer.is_linked_to_previous = False

        add_heading(doc, subtopics[0])
        add_text(doc, texts[0])
        doc.add_page_break()

        add_heading(doc, "ГЛАВА 1. " + subtopics[1])
        add_subheading(doc, "1.1. " + subtopics[2])
        add_text(doc, texts[1])
        doc.add_page_break()

        add_subheading(doc, "1.2. " + subtopics[3])
        add_text(doc, texts[2])
        doc.add_page_break()

        add_subheading(doc, "1.3. " + subtopics[4])
        add_text(doc, texts[3])
        doc.add_page_break()

        add_heading(doc, "ГЛАВА 2. " + subtopics[5])
        add_subheading(doc, "2.1. " + subtopics[6])
        add_text(doc, texts[4])
        doc.add_page_break()

        add_subheading(doc, "2.2. " + subtopics[7])
        add_text(doc, texts[5])
        doc.add_page_break()

        add_subheading(doc, "2.3. " + subtopics[8])
        add_text(doc, texts[6])
        doc.add_page_break()

        add_heading(doc, subtopics[9])
        add_text(doc, texts[7])
        doc.add_page_break()

        add_heading(doc, subtopics[10])

        bibliography = texts[8].split("\n")
        for i, reference in enumerate(bibliography, start=1):
            doc.add_paragraph(f"{i}. {reference.strip()}")

        # Добавляем номера страниц
        add_page_number(doc.sections[1])

        return save_doc_in_media(doc, sanitized_theme)


    # Если язык работы английский
    if language_of_work == "3":
        DEFAULT_FONT_NAME = "Times New Roman"

        # Функция для создания элемента
        def create_element(name):
            return OxmlElement(name)

        # Функция для создания атрибута
        def create_attribute(element, name, value):
            element.set(ns.qn(name), value)

        # Функция для добавления номера страницы
        def add_page_number(section):
            p = section.footer.paragraphs[0]  # Получаем параграф из нижнего колонтитула
            p.alignment = 1  # Выравнивание по центру
            p.style.font.size = Pt(12)  # Размер шрифта
            p.style.font.name = DEFAULT_FONT_NAME  # Шрифт

            # Создаем элементы XML для номера страницы
            run = p.add_run()
            fldChar1 = create_element('w:fldChar')
            create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = create_element('w:instrText')
            create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = create_element('w:fldChar')
            create_attribute(fldChar2, 'w:fldCharType', 'end')

            # Добавляем элементы в параграф для форматирования номера страницы
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        def add_heading(document, text, level=1):
            # Добавляем заголовок
            heading = document.add_heading(text, level=level)
            heading.bold = True
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.text = run.text.upper()  # Преобразуем текст заголовка в верхний регистр
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if text == "Introduction" or text == "Conclusion" or text == "References":
                heading.paragraph_format.line_spacing = 2  # Задаем межстрочный интервал
            else:
                heading.paragraph_format.line_spacing = 1  # Задаем межстрочный интервал

        def add_subheading(document, text, level=2):
            # Добавляем подзаголовок
            heading = document.add_heading(level=level)
            heading_run = heading.add_run(text)
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.paragraph_format.line_spacing = 2  # Задаем межстрочный интервал

        def add_text(document, text):
            # Добавляем текст
            style = document.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            paragraphs = text.split("\n")
            for paragraph_text in paragraphs:
                paragraph = document.add_paragraph()
                paragraph.add_run("\t" + paragraph_text)  # Добавляем текст с отступом
                paragraph.paragraph_format.right_indent = docx.shared.Inches(-0.590551)  # Устанавливаем правый отступ

        doc = Document()

        doc.sections[0].top_margin = docx.shared.Cm(2)
        # Если надо добавить титульный лист
        if cover_page_data == "1" or cover_page_data == "3":
            # Добавляем текст "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КЫРГЫЗСКОЙ РЕСПУБЛИКИ"
            ministry_title = doc.add_paragraph("MINISTRY OF EDUCATION AND SCIENCE OF THE KYRGYZ REPUBLIC")
            ministry_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            ministry_title.runs[0].font.name = "Times New Roman"
            ministry_title.runs[0].font.size = Pt(14)
            ministry_title.paragraph_format.line_spacing = 1.15

            # Добавляем название университета
            uni_title = doc.add_paragraph(university)
            uni_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            uni_title.runs[0].font.name = "Times New Roman"
            uni_title.runs[0].font.size = Pt(14)
            uni_title.runs[0].text = uni_title.runs[0].text.upper()
            uni_title.paragraph_format.line_spacing = 1.15

            if work_type == "Student's independent work":
                # Добавляем пустые строки
                for _ in range(3):
                    doc.add_paragraph()
            else:
                # Добавляем пустые строки
                for _ in range(4):
                    doc.add_paragraph()

            # Добавляем тип работы
            type_works_title = doc.add_paragraph(work_type)
            type_works_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            type_works_title.runs[0].font.name = "Times New Roman"
            type_works_title.runs[0].font.size = Pt(36)
            type_works_title.runs[0].bold = True
            type_works_title.paragraph_format.line_spacing = 1.15

            # Добавляем название темы
            theme_title = doc.add_paragraph("On the topic: " + work_theme)
            theme_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            theme_title.runs[0].font.name = "Times New Roman"
            theme_title.runs[0].font.size = Pt(14)
            theme_title.paragraph_format.line_spacing = 1.15

            if work_type == "Student's independent work":
                # Проверяем длину темы
                if len(work_theme) < 59:
                    # Добавляем пустые строки
                    for _ in range(4):
                        doc.add_paragraph()
                elif 59 <= len(work_theme) < 113:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 114 <= len(work_theme) < 179:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 180 <= len(work_theme):
                    # Добавляем пустые строки
                    for _ in range(2):
                        doc.add_paragraph()
            else:
                # Проверяем длину темы
                if len(work_theme) < 113:
                    # Добавляем пустые строки
                    for _ in range(4):
                        doc.add_paragraph()
                elif 114 <= len(work_theme) < 179:
                    # Добавляем пустые строки
                    for _ in range(3):
                        doc.add_paragraph()
                elif 180 <= len(work_theme):
                    # Добавляем пустые строки
                    for _ in range(1):
                        doc.add_paragraph()

            # Добавляем текст "Выполнил(а):" ФИО
            author_title = doc.add_paragraph("Completed: " + author_name)
            author_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            author_title.runs[0].font.name = "Times New Roman"
            author_title.runs[0].font.size = Pt(14)
            author_title.paragraph_format.line_spacing = 1.15

            # Добавляем текст "Группа:" название группы
            group_title = doc.add_paragraph("Group: " + group_name)
            group_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            group_title.runs[0].font.name = "Times New Roman"
            group_title.runs[0].font.size = Pt(14)
            group_title.paragraph_format.line_spacing = 1.15

            # Добавляем текст "Проверил(а):" ФИО преподавателя
            teacher_title = doc.add_paragraph("Checked: " + teacher_name)
            teacher_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            teacher_title.runs[0].font.name = "Times New Roman"
            teacher_title.runs[0].font.size = Pt(14)
            teacher_title.paragraph_format.line_spacing = 1.15

            if len(university) <= 89:
                # Добавляем пустые строки
                for _ in range(3):
                    doc.add_paragraph()
            elif len(university) > 89:
                # Добавляем пустые строки
                for _ in range(2):
                    doc.add_paragraph()

            # Получаем текущую дату и извлекаем год
            current_year = datetime.datetime.now().year

            # Добавляем текст "Бишкек -" текущий год
            centered_paragraph = doc.add_paragraph()
            centered_run = centered_paragraph.add_run("Bishkek - " + str(current_year))
            centered_run.font.name = "Times New Roman"
            centered_run.font.size = Pt(14)
            centered_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            centered_paragraph.paragraph_format.line_spacing = 1.15

            doc.add_paragraph()

        # Добавляем заголовок для оглавления
        toc_title = doc.add_paragraph("CONTENTS")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in toc_title.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True

        # Добавляем пустой параграф перед оглавлением
        toc_empty_para = doc.add_paragraph()
        toc_empty_para.paragraph_format.space_before = Inches(0)
        toc_empty_para.paragraph_format.space_after = Inches(0)
        toc_run = toc_empty_para.add_run()

        # Создаем XML-элементы для оглавления
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # измените "1-3" в зависимости от необходимого уровня заголовков

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar3 = OxmlElement('w:instrText')  # Заменяем на 'w:instrText', чтобы добавить текст инструкции
        fldChar3.set(qn('xml:space'), 'preserve')
        fldChar3.text = "Right-click to update field."

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        # Добавляем XML-элементы в оглавление
        toc_run._r.append(fldChar)
        toc_run._r.append(instrText)
        toc_run._r.append(fldChar2)
        toc_run._r.append(fldChar3)  # Добавляем элемент с текстом инструкции
        toc_run._r.append(fldChar4)

        # Добавляем разрыв страницы
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Добавляем раздел, из которого мы хотим получить номера страниц
        doc.add_section(WD_SECTION.CONTINUOUS)

        # Добавляем текст на первой странице
        doc.sections[1].footer.is_linked_to_previous = False

        add_heading(doc, subtopics[0])
        add_text(doc, texts[0])
        doc.add_page_break()

        add_heading(doc, "CHAPTER 1. " + subtopics[1])
        add_subheading(doc, "1.1. " + subtopics[2])
        add_text(doc, texts[1])
        doc.add_page_break()

        add_subheading(doc, "1.2. " + subtopics[3])
        add_text(doc, texts[2])
        doc.add_page_break()

        add_subheading(doc, "1.3. " + subtopics[4])
        add_text(doc, texts[3])
        doc.add_page_break()

        add_heading(doc, "CHAPTER 2. " + subtopics[5])
        add_subheading(doc, "2.1. " + subtopics[6])
        add_text(doc, texts[4])
        doc.add_page_break()

        add_subheading(doc, "2.2. " + subtopics[7])
        add_text(doc, texts[5])
        doc.add_page_break()

        add_subheading(doc, "2.3. " + subtopics[8])
        add_text(doc, texts[6])
        doc.add_page_break()

        add_heading(doc, subtopics[9])
        add_text(doc, texts[7])
        doc.add_page_break()

        add_heading(doc, subtopics[10])
        bibliography = texts[8].split("\n")
        for i, reference in enumerate(bibliography, start=1):
            doc.add_paragraph(f"{i}. {reference.strip()}")

        # Добавляем номера страниц
        add_page_number(doc.sections[1])

        return save_doc_in_media(doc, sanitized_theme)
